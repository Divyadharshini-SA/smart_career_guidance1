import re
import os

# Try importing NLP libs gracefully
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_SUPPORT = True
except ImportError:
    SKLEARN_SUPPORT = False

try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    nltk.download('punkt',     quiet=True)
    nltk.download('stopwords', quiet=True)
    NLTK_SUPPORT = True
except ImportError:
    NLTK_SUPPORT = False


SKILL_KEYWORDS = [
    # Programming
    "python", "java", "c++", "c", "javascript", "typescript", "golang", "rust", "kotlin", "swift",
    # Web
    "html", "css", "react", "angular", "vue", "node", "express", "django", "flask", "fastapi",
    # Data / ML
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras", "scikit-learn",
    "pandas", "numpy", "matplotlib", "data science", "nlp", "computer vision",
    # Cloud / DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "ci/cd", "terraform",
    # DB
    "mysql", "postgresql", "mongodb", "redis", "sqlite", "firebase", "elasticsearch",
    # Tools
    "git", "linux", "bash", "rest api", "graphql", "microservices", "agile", "scrum",
    # Soft skills
    "communication", "leadership", "teamwork", "problem solving", "critical thinking",
]

INDUSTRY_REQUIRED = {
    "software_engineer": ["python","java","data structures","algorithms","git","sql","rest api","agile"],
    "data_scientist"   : ["python","machine learning","pandas","numpy","statistics","sql","data visualization"],
    "web_developer"    : ["html","css","javascript","react","node","git","rest api"],
    "ml_engineer"      : ["python","machine learning","deep learning","tensorflow","pytorch","mlops","docker"],
}


class ResumeService:

    def extract_text(self, filepath: str) -> str:
        ext = os.path.splitext(filepath)[1].lower()
        text = ""
        try:
            if ext == '.pdf' and PDF_SUPPORT:
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += (page.extract_text() or "") + "\n"
            elif ext in ('.doc', '.docx') and DOCX_SUPPORT:
                doc = Document(filepath)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                # Fallback: read as plain text
                with open(filepath, 'r', errors='ignore') as f:
                    text = f.read()
        except Exception as e:
            text = f"[Error extracting text: {e}]"
        return text.lower()

    def extract_skills(self, text: str) -> list:
        found = []
        for skill in SKILL_KEYWORDS:
            if skill in text:
                found.append(skill)
        return found

    def calculate_resume_score(self, text: str, extracted_skills: list) -> float:
        """
        Score based on:
          - Skill coverage      (50%)
          - Resume sections     (30%)
          - Text quality length (20%)
        """
        # Skill coverage score
        skill_score = min(len(extracted_skills) / 15, 1.0) * 50

        # Section detection
        sections = ['experience', 'education', 'skills', 'projects', 'achievements', 'certifications']
        found_sections = sum(1 for s in sections if s in text)
        section_score = (found_sections / len(sections)) * 30

        # Length score
        words = len(text.split())
        length_score = min(words / 400, 1.0) * 20

        total = float(f"{skill_score + section_score + length_score:.2f}")
        return total


    def generate_feedback(self, text: str, skills: list, score: float) -> dict:
        """Generate detailed actionable feedback on the resume."""

        # ── Section checks ──────────────────────────────────────
        sections_check = {
            'Contact Info'    : any(k in text for k in ['email','phone','linkedin','github','contact']),
            'Summary/Objective': any(k in text for k in ['summary','objective','about','profile']),
            'Skills Section'  : any(k in text for k in ['skills','technologies','tools']),
            'Projects'        : any(k in text for k in ['project','built','developed','created']),
            'Education'       : any(k in text for k in ['education','university','college','degree','b.tech','btech','b.e']),
            'Experience'      : any(k in text for k in ['experience','internship','worked','intern','company']),
            'Certifications'  : any(k in text for k in ['certification','certified','certificate','coursera','udemy','nptel']),
            'Achievements'    : any(k in text for k in ['achievement','award','winner','rank','topper','merit']),
        }

        # ── Strengths ───────────────────────────────────────────
        strengths = []
        if len(skills) >= 10:
            strengths.append(f'Strong skill set with {len(skills)} detected technical skills')
        elif len(skills) >= 5:
            strengths.append(f'Good skill coverage with {len(skills)} skills listed')
        if sections_check['Projects']:
            strengths.append('Projects section present — shows practical experience')
        if sections_check['Experience']:
            strengths.append('Work/internship experience mentioned — great for placements')
        if sections_check['Certifications']:
            strengths.append('Certifications listed — adds credibility')
        if sections_check['Achievements']:
            strengths.append('Achievements/awards section found — stands out to recruiters')
        if len(text.split()) >= 400:
            strengths.append('Resume has good content length (400+ words)')

        if not strengths:
            strengths.append('Resume uploaded successfully — review suggestions below to improve')

        # ── Missing sections → improvements ─────────────────────
        improvements = []
        if not sections_check['Summary/Objective']:
            improvements.append('Add a 2-3 line Professional Summary at the top highlighting your goals')
        if not sections_check['Skills Section']:
            improvements.append('Add a dedicated Skills section listing all technical tools and languages')
        if not sections_check['Projects']:
            improvements.append('Add 2-3 Projects with tech stack, GitHub link, and what problem you solved')
        if not sections_check['Certifications']:
            improvements.append('Add certifications (Coursera, NPTEL, Udemy) to boost your profile score')
        if not sections_check['Achievements']:
            improvements.append('Add academic achievements, hackathon wins, or contest rankings')
        if not sections_check['Experience']:
            improvements.append('Add internship or freelance experience — even personal projects count')
        if len(skills) < 5:
            improvements.append('List more technical skills — include all languages, tools, and frameworks you know')
        if len(text.split()) < 250:
            improvements.append('Resume seems short — add more details to projects and experience sections')

        # ── ATS keyword tips ────────────────────────────────────
        ats_tips = []
        common_ats_keywords = ['data structures', 'algorithms', 'rest api', 'agile', 'git', 'sql',
                                'python', 'java', 'javascript', 'react', 'node', 'machine learning']
        missing_ats_full = [k for k in common_ats_keywords if k not in text]
        missing_ats = [missing_ats_full[i] for i in range(min(4, len(missing_ats_full)))]
        if missing_ats:
            ats_tips.append(f'Add ATS keywords: {", ".join(missing_ats).title()}')
        ats_tips.append('Use standard section headings (Skills, Projects, Education, Experience)')
        ats_tips.append('Avoid tables, images, or columns — ATS cannot parse them')
        ats_tips.append('Use action verbs: Developed, Built, Designed, Implemented, Optimized')

        # ── Score label ─────────────────────────────────────────
        if score >= 75:
            score_label, score_color = 'Excellent Resume ⭐', '#06D6A0'
        elif score >= 60:
            score_label, score_color = 'Good Resume 👍', '#7C5CFC'
        elif score >= 45:
            score_label, score_color = 'Average Resume 📝', '#FFD93D'
        else:
            score_label, score_color = 'Needs Improvement ⚠️', '#FF6B6B'

        return {
            'score'        : score,
            'score_label'  : score_label,
            'score_color'  : score_color,
            'skill_count'  : len(skills),
            'sections'     : sections_check,
            'strengths'    : strengths,
            'improvements' : improvements,
            'ats_tips'     : ats_tips,
        }

    def tfidf_match(self, resume_text: str, job_description: str) -> float:
        """Cosine similarity between resume and job description using TF-IDF."""
        if not SKLEARN_SUPPORT:
            return 0.0
        try:
            vectorizer = TfidfVectorizer(stop_words='english')
            tfidf = vectorizer.fit_transform([resume_text, job_description])
            sim = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
            return float(f"{float(sim) * 100:.2f}")
        except Exception:
            return 0.0
        

# import re
# import os

# # Try importing NLP libs gracefully
# try:
#     import pdfplumber
#     PDF_SUPPORT = True
# except ImportError:
#     PDF_SUPPORT = False

# try:
#     from docx import Document
#     DOCX_SUPPORT = True
# except ImportError:
#     DOCX_SUPPORT = False

# try:
#     from sklearn.feature_extraction.text import TfidfVectorizer
#     from sklearn.metrics.pairwise import cosine_similarity
#     SKLEARN_SUPPORT = True
# except ImportError:
#     SKLEARN_SUPPORT = False

# try:
#     import nltk
#     from nltk.corpus import stopwords
#     from nltk.tokenize import word_tokenize
#     nltk.download('punkt',     quiet=True)
#     nltk.download('stopwords', quiet=True)
#     NLTK_SUPPORT = True
# except ImportError:
#     NLTK_SUPPORT = False


# SKILL_KEYWORDS = [
#     # Programming
#     "python", "java", "c++", "c", "javascript", "typescript", "golang", "rust", "kotlin", "swift",
#     # Web
#     "html", "css", "react", "angular", "vue", "node", "express", "django", "flask", "fastapi",
#     # Data / ML
#     "machine learning", "deep learning", "tensorflow", "pytorch", "keras", "scikit-learn",
#     "pandas", "numpy", "matplotlib", "data science", "nlp", "computer vision",
#     # Cloud / DevOps
#     "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "ci/cd", "terraform",
#     # DB
#     "mysql", "postgresql", "mongodb", "redis", "sqlite", "firebase", "elasticsearch",
#     # Tools
#     "git", "linux", "bash", "rest api", "graphql", "microservices", "agile", "scrum",
#     # Soft skills
#     "communication", "leadership", "teamwork", "problem solving", "critical thinking",
# ]

# INDUSTRY_REQUIRED = {
#     "software_engineer": ["python","java","data structures","algorithms","git","sql","rest api","agile"],
#     "data_scientist"   : ["python","machine learning","pandas","numpy","statistics","sql","data visualization"],
#     "web_developer"    : ["html","css","javascript","react","node","git","rest api"],
#     "ml_engineer"      : ["python","machine learning","deep learning","tensorflow","pytorch","mlops","docker"],
# }


# class ResumeService:

#     def extract_text(self, filepath: str) -> str:
#         ext = os.path.splitext(filepath)[1].lower()
#         text = ""
#         try:
#             if ext == '.pdf' and PDF_SUPPORT:
#                 with pdfplumber.open(filepath) as pdf:
#                     for page in pdf.pages:
#                         text += (page.extract_text() or "") + "\n"
#             elif ext in ('.doc', '.docx') and DOCX_SUPPORT:
#                 doc = Document(filepath)
#                 text = "\n".join([p.text for p in doc.paragraphs])
#             else:
#                 # Fallback: read as plain text
#                 with open(filepath, 'r', errors='ignore') as f:
#                     text = f.read()
#         except Exception as e:
#             text = f"[Error extracting text: {e}]"
#         return text.lower()

#     def extract_skills(self, text: str) -> list:
#         found = []
#         for skill in SKILL_KEYWORDS:
#             if skill in text:
#                 found.append(skill)
#         return found

#     def calculate_resume_score(self, text: str, extracted_skills: list) -> float:
#         """
#         Score based on:
#           - Skill coverage      (50%)
#           - Resume sections     (30%)
#           - Text quality length (20%)
#         """
#         # Skill coverage score
#         skill_score = min(len(extracted_skills) / 15, 1.0) * 50

#         # Section detection
#         sections = ['experience', 'education', 'skills', 'projects', 'achievements', 'certifications']
#         found_sections = sum(1 for s in sections if s in text)
#         section_score = (found_sections / len(sections)) * 30

#         # Length score
#         words = len(text.split())
#         length_score = min(words / 400, 1.0) * 20

#         total = round(skill_score + section_score + length_score, 2)
#         return total

#     def tfidf_match(self, resume_text: str, job_description: str) -> float:
#         """Cosine similarity between resume and job description using TF-IDF."""
#         if not SKLEARN_SUPPORT:
#             return 0.0
#         try:
#             vectorizer = TfidfVectorizer(stop_words='english')
#             tfidf = vectorizer.fit_transform([resume_text, job_description])
#             sim = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
#             return round(float(sim) * 100, 2)
#         except Exception:
#             return 0.0